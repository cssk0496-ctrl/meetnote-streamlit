import json
import csv
import html
from io import BytesIO, StringIO
from datetime import datetime

import streamlit as st
import streamlit.components.v1 as components
from openai import OpenAI
from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
from reportlab.lib.styles import ParagraphStyle


st.set_page_config(
    page_title="MeetNote | AI 회의록",
    page_icon="🎙️",
    layout="wide",
    initial_sidebar_state="collapsed",
)

st.markdown(
    """
    <style>
    .stApp {
        background:
          radial-gradient(circle at 14% 0%, rgba(99, 82, 225, .20), transparent 31rem),
          radial-gradient(circle at 88% 28%, rgba(38, 116, 255, .12), transparent 28rem),
          #080d19;
        color: #f8fafc;
    }
    [data-testid="stHeader"] { background: transparent; }
    [data-testid="stSidebar"] { background: #0c1220; }
    [data-testid="stAlert"] {
        color: #dbe7ff !important;
        background: rgba(53, 72, 112, .45) !important;
        border: 1px solid rgba(121, 155, 226, .28) !important;
    }
    [data-testid="stAlert"] * { color: #dbe7ff !important; }
    .block-container { max-width: 1160px; padding-top: 1.6rem; padding-bottom: 5rem; }
    .brand {
        display: flex; align-items: center; gap: 11px; padding: 4px 0 22px;
        border-bottom: 1px solid rgba(154,169,201,.14);
        font-size: 20px; font-weight: 800; letter-spacing: -.04em;
    }
    .brand-mark {
        display: grid; width: 36px; height: 36px; place-items: center;
        border-radius: 11px; background: linear-gradient(145deg,#7c5cff,#395bd8);
        box-shadow: 0 8px 28px rgba(91,79,223,.35);
    }
    .hero { padding: 76px 0 28px; }
    .eyebrow { color:#aaa0ff; font-size:13px; font-weight:800; letter-spacing:.08em; }
    .hero h1 {
        margin: 18px 0 22px; font-size: clamp(46px,5.2vw,72px);
        line-height:1.08; letter-spacing:-.065em;
    }
    .hero h1 span {
        color:transparent; background:linear-gradient(90deg,#a283ff,#5da8ff,#72e4ff);
        -webkit-background-clip:text; background-clip:text;
    }
    .hero p { color:#9ca8bd; font-size:17px; line-height:1.75; }
    .feature-line { display:flex; gap:22px; margin-top:28px; color:#77849a; font-size:12px; }
    .upload-shell {
        margin-top: 36px; padding: 28px; border:1px solid rgba(151,164,200,.20);
        border-radius:24px; background:linear-gradient(150deg,rgba(26,35,59,.95),rgba(13,20,36,.93));
        box-shadow:0 30px 80px rgba(0,0,0,.35);
    }
    .upload-shell h3 { margin:2px 0 4px; }
    [data-testid="stFileUploaderDropzone"] {
        min-height:170px; border:1px dashed rgba(139,92,246,.55);
        border-radius:18px; background:rgba(74,71,150,.08);
    }
    [data-testid="stFileUploader"] label,
    [data-testid="stFileUploader"] label p {
        color:#9ca8bd !important;
    }
    [data-testid="stFileUploaderDropzoneInstructions"] *,
    [data-testid="stFileUploaderDropzone"] small {
        color:#e3e6ee !important; font-weight:700;
    }
    [data-testid="stFileUploaderDropzone"] button {
        color:#fff !important;
        background:#4c5fc2 !important;
        border:1px solid rgba(255,255,255,.22) !important;
    }
    [data-testid="stFileUploaderDropzone"] button * {
        color:#fff !important;
    }
    .stButton > button, .stDownloadButton > button,
    div[data-testid="stButton"] > button,
    div[data-testid="stDownloadButton"] > button {
        width:100%; min-height:48px; border:0; border-radius:13px;
        color:#fff !important;
        background:linear-gradient(100deg,#7857ef,#4277ef) !important;
        font-weight:800; box-shadow:0 12px 30px rgba(79,70,229,.22);
    }
    .stButton > button *,
    .stDownloadButton > button *,
    div[data-testid="stButton"] > button *,
    div[data-testid="stDownloadButton"] > button * {
        color:#fff !important;
    }
    .stButton > button:hover, .stDownloadButton > button:hover {
        border:0; color:#fff !important; filter:brightness(1.08);
    }
    .stButton > button:disabled,
    div[data-testid="stButton"] > button:disabled {
        color:rgba(255,255,255,.68) !important;
        background:linear-gradient(100deg,#4f3e94,#31538f) !important;
        opacity:.72 !important;
    }
    .stButton > button:disabled *,
    div[data-testid="stButton"] > button:disabled * {
        color:rgba(255,255,255,.68) !important;
    }
    .result-title { margin:72px 0 22px; }
    .done-label { color:#69e59a; font-size:12px; font-weight:800; }
    .card {
        height:100%; padding:25px; border:1px solid rgba(154,169,201,.16);
        border-radius:18px; background:rgba(17,24,42,.72);
    }
    .card h3 { margin:0 0 17px; font-size:16px; }
    .card p, .card li { color:#a5b0c3; font-size:13px; line-height:1.8; }
    .action {
        margin:10px 0; padding:13px 14px; border-radius:11px;
        background:rgba(6,11,21,.46); color:#dfe4ee; font-size:12px;
    }
    .action small { display:block; margin-top:5px; color:#748198; }
    .transcript {
        padding:22px; border:1px solid rgba(154,169,201,.16);
        border-radius:16px; background:rgba(17,24,42,.66);
        color:#a5b0c3; font-size:13px; line-height:1.9; white-space:pre-wrap;
    }
    .security {
        margin-top:60px; padding:25px; border:1px solid rgba(69,177,255,.14);
        border-radius:20px; background:rgba(38,54,92,.28);
    }
    .security strong { color:#dbe7ff; }
    .security p { margin:8px 0 0; color:#7d8ba2; font-size:12px; }
    @media (max-width: 640px) {
        .block-container { padding-left:1rem; padding-right:1rem; }
        .hero { padding-top:48px; }
        .hero h1 { font-size:44px; }
        .feature-line { flex-direction:column; gap:8px; }
        .upload-shell { padding:18px; }
    }
    </style>
    """,
    unsafe_allow_html=True,
)


SUMMARY_SCHEMA = {
    "type": "object",
    "properties": {
        "title": {"type": "string"},
        "summary": {"type": "string"},
        "decisions": {"type": "array", "items": {"type": "string"}},
        "actions": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "task": {"type": "string"},
                    "owner": {"type": "string"},
                    "due": {"type": "string"},
                },
                "required": ["task", "owner", "due"],
                "additionalProperties": False,
            },
        },
        "keywords": {"type": "array", "items": {"type": "string"}},
    },
    "required": ["title", "summary", "decisions", "actions", "keywords"],
    "additionalProperties": False,
}

SAMPLE_RESULT = {
    "title": "주간 프로젝트 진행회의",
    "summary": (
        "AI 회의록 자동화 서비스의 주요 기능과 적용 일정을 검토했습니다. "
        "음성 파일에서 핵심 내용과 실행 업무를 자동 추출하고, 결과를 Word·PDF·텍스트로 "
        "내보낼 수 있도록 구성하기로 했습니다. 다음 회의 전까지 테스트 결과를 취합합니다."
    ),
    "decisions": [
        "회의록 결과에 핵심 요약, 결정사항, 담당 업무를 구분하여 표시합니다.",
        "Word, PDF, 텍스트 및 담당 업무 CSV 다운로드 기능을 적용합니다.",
        "짧은 음성 파일로 1차 기능 검증을 진행합니다.",
    ],
    "actions": [
        {"task": "테스트 음성 파일 준비", "owner": "김창숙", "due": "8월 3일"},
        {"task": "회의록 다운로드 결과 확인", "owner": "기획팀", "due": "8월 4일"},
        {"task": "개선 의견 취합", "owner": "참석자 전체", "due": "다음 회의 전"},
    ],
    "keywords": ["AI 회의록", "음성인식", "업무 자동화", "결과 내보내기"],
}

SAMPLE_TRANSCRIPT = """[00:00] 김창숙: 오늘은 AI 회의록 자동화 서비스의 주요 기능을 검토하겠습니다.
[00:12] 참석자 1: 요약뿐 아니라 결정사항과 담당 업무도 구분되면 좋겠습니다.
[00:25] 김창숙: Word, PDF, 텍스트 파일과 담당 업무 CSV를 내려받을 수 있게 적용하겠습니다.
[00:41] 참석자 2: 다음 회의 전까지 짧은 음성 파일로 기능을 검증하겠습니다."""


def get_api_key() -> str | None:
    try:
        return st.secrets["OPENAI_API_KEY"]
    except (KeyError, FileNotFoundError):
        return None


def analyze_meeting(uploaded_file, language: str) -> tuple[str, dict]:
    client = OpenAI(api_key=get_api_key())

    transcription = client.audio.transcriptions.create(
        model="gpt-transcribe",
        file=(uploaded_file.name, uploaded_file.getvalue(), uploaded_file.type),
        prompt=(
            f"이 파일은 {language} 업무 회의입니다. "
            "회사명, 프로젝트명, 사람 이름과 기술 용어를 가능한 정확히 기록하세요."
        ),
    )
    transcript = transcription.text

    response = client.responses.create(
        model="gpt-5-mini",
        instructions=(
            "당신은 한국 기업의 전문 회의록 작성자입니다. "
            "녹취에 없는 내용을 추측하지 말고, 불명확한 담당자나 기한은 '미정'으로 표시하세요. "
            "결정된 내용과 단순 의견을 구분하고, 실행 업무는 짧고 명확하게 작성하세요."
        ),
        input=f"다음 회의 녹취를 한국어 회의록으로 정리하세요.\n\n{transcript}",
        text={
            "format": {
                "type": "json_schema",
                "name": "meeting_minutes",
                "strict": True,
                "schema": SUMMARY_SCHEMA,
            }
        },
    )
    return transcript, json.loads(response.output_text)


def make_minutes_text(result: dict, transcript: str, filename: str) -> str:
    lines = [
        f"# {result['title']}",
        f"- 원본 파일: {filename}",
        f"- 작성일: {datetime.now().strftime('%Y-%m-%d %H:%M')}",
        "",
        "## 핵심 요약",
        result["summary"],
        "",
        "## 결정사항",
    ]
    lines.extend(
        [f"{index}. {item}" for index, item in enumerate(result["decisions"], 1)]
        or ["- 확정된 결정사항 없음"]
    )
    lines.extend(["", "## 담당 업무"])
    lines.extend(
        [
            f"- {item['task']} / 담당: {item['owner']} / 기한: {item['due']}"
            for item in result["actions"]
        ]
        or ["- 등록된 담당 업무 없음"]
    )
    lines.extend(["", "## 키워드", ", ".join(result["keywords"]), "", "## 전체 녹취", transcript])
    return "\n".join(lines)


def make_word_file(result: dict, transcript: str, filename: str) -> bytes:
    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal.font.size = Pt(10)

    document.add_heading(result["title"], level=0)
    document.add_paragraph(f"원본 파일: {filename}")
    document.add_paragraph(f"작성일: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    document.add_heading("핵심 요약", level=1)
    document.add_paragraph(result["summary"])

    document.add_heading("결정사항", level=1)
    if result["decisions"]:
        for item in result["decisions"]:
            document.add_paragraph(item, style="List Number")
    else:
        document.add_paragraph("확정된 결정사항 없음")

    document.add_heading("담당 업무", level=1)
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for cell, title in zip(table.rows[0].cells, ["담당 업무", "담당자", "기한"]):
        cell.text = title
    for item in result["actions"]:
        cells = table.add_row().cells
        cells[0].text = item["task"]
        cells[1].text = item["owner"]
        cells[2].text = item["due"]

    document.add_heading("주요 키워드", level=1)
    document.add_paragraph(", ".join(result["keywords"]))
    document.add_heading("전체 녹취", level=1)
    document.add_paragraph(transcript)

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def make_pdf_file(result: dict, transcript: str, filename: str) -> bytes:
    output = BytesIO()
    pdfmetrics.registerFont(UnicodeCIDFont("HYSMyeongJo-Medium"))
    document = SimpleDocTemplate(
        output,
        pagesize=A4,
        rightMargin=42,
        leftMargin=42,
        topMargin=45,
        bottomMargin=45,
        title=result["title"],
    )
    title_style = ParagraphStyle(
        "KoreanTitle",
        fontName="HYSMyeongJo-Medium",
        fontSize=19,
        leading=27,
        spaceAfter=16,
    )
    heading_style = ParagraphStyle(
        "KoreanHeading",
        fontName="HYSMyeongJo-Medium",
        fontSize=13,
        leading=19,
        spaceBefore=14,
        spaceAfter=7,
    )
    body_style = ParagraphStyle(
        "KoreanBody",
        fontName="HYSMyeongJo-Medium",
        fontSize=9.5,
        leading=16,
        spaceAfter=5,
    )

    def safe(value: str) -> str:
        return html.escape(str(value)).replace("\n", "<br/>")

    story = [
        Paragraph(safe(result["title"]), title_style),
        Paragraph(f"원본 파일: {safe(filename)}", body_style),
        Paragraph(f"작성일: {datetime.now().strftime('%Y-%m-%d %H:%M')}", body_style),
        Spacer(1, 8),
        Paragraph("핵심 요약", heading_style),
        Paragraph(safe(result["summary"]), body_style),
        Paragraph("결정사항", heading_style),
    ]
    decisions = result["decisions"] or ["확정된 결정사항 없음"]
    story.extend(
        Paragraph(f"{index}. {safe(item)}", body_style)
        for index, item in enumerate(decisions, 1)
    )
    story.append(Paragraph("담당 업무", heading_style))
    actions = result["actions"] or [{"task": "등록된 담당 업무 없음", "owner": "-", "due": "-"}]
    story.extend(
        Paragraph(
            f"• {safe(item['task'])} / 담당: {safe(item['owner'])} / 기한: {safe(item['due'])}",
            body_style,
        )
        for item in actions
    )
    story.extend(
        [
            Paragraph("주요 키워드", heading_style),
            Paragraph(safe(", ".join(result["keywords"])), body_style),
            Paragraph("전체 녹취", heading_style),
            Paragraph(safe(transcript), body_style),
        ]
    )
    document.build(story)
    return output.getvalue()


def make_actions_csv(result: dict) -> bytes:
    output = StringIO()
    writer = csv.writer(output)
    writer.writerow(["담당 업무", "담당자", "기한"])
    for item in result["actions"]:
        writer.writerow([item["task"], item["owner"], item["due"]])
    return output.getvalue().encode("utf-8-sig")


def show_copy_button(minutes: str) -> None:
    payload = json.dumps(minutes)
    components.html(
        f"""
        <button id="copyMinutes" style="
          width:100%;height:46px;border:1px solid rgba(139,92,246,.45);
          border-radius:12px;color:#fff;background:linear-gradient(100deg,#7857ef,#4277ef);
          font:700 14px sans-serif;cursor:pointer;">
          📋 전체 회의록 복사하기
        </button>
        <script>
          const button = document.getElementById("copyMinutes");
          button.addEventListener("click", async () => {{
            try {{
              await navigator.clipboard.writeText({payload});
              button.textContent = "✓ 복사 완료";
              setTimeout(() => button.textContent = "📋 전체 회의록 복사하기", 1800);
            }} catch (error) {{
              button.textContent = "복사 권한을 확인해 주세요";
            }}
          }});
        </script>
        """,
        height=52,
    )


st.markdown(
    '<div class="brand"><span class="brand-mark">✦</span>MeetNote</div>',
    unsafe_allow_html=True,
)

left, right = st.columns([1.02, 0.98], gap="large", vertical_alignment="center")

with left:
    st.markdown(
        """
        <div class="hero">
          <div class="eyebrow">✦ AI 회의 업무 자동화</div>
          <h1>회의는 말로,<br><span>정리는 AI에게.</span></h1>
          <p>녹음 파일을 올리면 발언 내용을 텍스트로 변환하고,<br>
          핵심 요약부터 결정사항과 담당 업무까지 한 번에 정리합니다.</p>
          <div class="feature-line">
            <span>◷ 몇 분 안에 결과 확인</span>
            <span>♢ 안전한 파일 처리</span>
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

with right:
    st.markdown('<div class="upload-shell">', unsafe_allow_html=True)
    st.markdown("#### 🎙️ 회의 파일 업로드")
    language = st.selectbox(
        "회의 언어",
        ["한국어", "영어", "한국어와 영어 혼용"],
        label_visibility="collapsed",
    )
    uploaded_file = st.file_uploader(
        "MP3, WAV, M4A 또는 MP4 파일을 올려주세요.",
        type=["mp3", "wav", "m4a", "mp4", "mpeg", "mpga", "ogg", "webm"],
        help="회의에 개인정보나 기밀정보가 있다면 사내 보안정책을 먼저 확인하세요.",
    )

    if not get_api_key():
        st.info("배포 후 Streamlit Secrets에 OPENAI_API_KEY를 등록하면 AI 분석이 활성화됩니다.")

    analyze = st.button(
        "✦ AI 회의록 만들기",
        type="primary",
        use_container_width=True,
        disabled=uploaded_file is None or not get_api_key(),
    )
    sample = st.button(
        "🧪 샘플 회의록 보기",
        use_container_width=True,
        help="음성 파일이나 API 사용 없이 결과 화면과 다운로드 기능을 확인합니다.",
    )
    st.caption("🔒 업로드한 파일은 앱에 영구 저장하지 않습니다.")
    st.markdown("</div>", unsafe_allow_html=True)

if sample:
    st.session_state["meeting_result"] = SAMPLE_RESULT
    st.session_state["meeting_transcript"] = SAMPLE_TRANSCRIPT
    st.session_state["meeting_filename"] = "샘플_회의음성.m4a"
    st.session_state["is_sample"] = True

if analyze and uploaded_file:
    progress = st.progress(8, text="회의 음성을 준비하고 있습니다...")
    try:
        progress.progress(28, text="음성을 텍스트로 변환하고 있습니다...")
        transcript, result = analyze_meeting(uploaded_file, language)
        progress.progress(78, text="핵심 내용과 담당 업무를 정리하고 있습니다...")
        st.session_state["meeting_result"] = result
        st.session_state["meeting_transcript"] = transcript
        st.session_state["meeting_filename"] = uploaded_file.name
        st.session_state["is_sample"] = False
        progress.progress(100, text="회의록이 완성되었습니다.")
    except Exception as error:
        st.error(f"분석 중 오류가 발생했습니다: {error}")
    finally:
        progress.empty()

if "meeting_result" in st.session_state:
    result = st.session_state["meeting_result"]
    transcript = st.session_state["meeting_transcript"]
    filename = st.session_state["meeting_filename"]
    minutes = make_minutes_text(result, transcript, filename)
    word_file = make_word_file(result, transcript, filename)
    pdf_file = make_pdf_file(result, transcript, filename)
    actions_csv = make_actions_csv(result)

    st.markdown(
        f"""
        <div class="result-title">
          <div class="done-label">✓ {"샘플 결과" if st.session_state.get("is_sample") else "분석 완료"}</div>
          <h2>{result['title']}</h2>
        </div>
        """,
        unsafe_allow_html=True,
    )

    st.markdown("##### 회의록 내보내기")
    word_col, pdf_col, text_col, csv_col = st.columns(4)
    with word_col:
        st.download_button(
            "📝 Word 회의록",
            data=word_file,
            file_name=f"{result['title']}_회의록.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )
    with pdf_col:
        st.download_button(
            "📄 PDF 회의록",
            data=pdf_file,
            file_name=f"{result['title']}_회의록.pdf",
            mime="application/pdf",
            use_container_width=True,
        )
    with text_col:
        st.download_button(
            "📃 텍스트 파일",
            data=minutes.encode("utf-8-sig"),
            file_name=f"{result['title']}_회의록.txt",
            mime="text/plain",
            use_container_width=True,
        )
    with csv_col:
        st.download_button(
            "📊 담당 업무 CSV",
            data=actions_csv,
            file_name=f"{result['title']}_담당업무.csv",
            mime="text/csv",
            use_container_width=True,
        )
    show_copy_button(minutes)

    summary_tab, transcript_tab = st.tabs(["✦ AI 요약", "☰ 전체 녹취"])
    with summary_tab:
        st.markdown(
            f'<div class="card"><h3>✦ 핵심 요약</h3><p>{result["summary"]}</p></div>',
            unsafe_allow_html=True,
        )
        decision_col, action_col = st.columns(2, gap="medium")
        with decision_col:
            decisions_html = "".join(f"<li>{item}</li>" for item in result["decisions"])
            if not decisions_html:
                decisions_html = "<li>확정된 결정사항이 없습니다.</li>"
            st.markdown(
                f'<div class="card"><h3>✓ 결정사항</h3><ul>{decisions_html}</ul></div>',
                unsafe_allow_html=True,
            )
        with action_col:
            actions_html = "".join(
                (
                    f'<div class="action"><strong>{item["task"]}</strong>'
                    f'<small>{item["owner"]} · {item["due"]}</small></div>'
                )
                for item in result["actions"]
            )
            if not actions_html:
                actions_html = '<div class="action">등록된 담당 업무가 없습니다.</div>'
            st.markdown(
                f'<div class="card"><h3>◎ 담당 업무</h3>{actions_html}</div>',
                unsafe_allow_html=True,
            )
        st.markdown("##### 주요 키워드")
        st.write(" · ".join(result["keywords"]))

    with transcript_tab:
        st.markdown(f'<div class="transcript">{transcript}</div>', unsafe_allow_html=True)

st.markdown(
    """
    <div class="security">
      <strong>♢ 회의 자료, 안전하게 다뤄야 하니까</strong>
      <p>API 키는 코드에 적지 않고 Streamlit Secrets에 저장하세요.
      민감한 사내 회의는 회사의 외부 AI 서비스 이용 기준을 먼저 확인해야 합니다.</p>
    </div>
    """,
    unsafe_allow_html=True,
)
